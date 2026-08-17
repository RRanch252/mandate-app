"""Build two sample documents used by the smoke test and for trying the app by hand.

The PDF is written from raw PDF syntax rather than with a rendering library so that
the repository needs no extra dependency just to produce a test file.

The facts below are invented for a fictional company. One fact is deliberately absent
(customer churn) so the smoke test can prove that Mandate says "not in the documents"
instead of producing a number.
"""

import os
import sys

PAGES = [
    [
        "ACME INDUSTRIAL HOLDINGS LIMITED",
        "Confidential information memorandum",
        "",
        "1. Business overview",
        "",
        "Acme Industrial Holdings Limited designs and manufactures precision hydraulic",
        "couplings for the offshore energy sector. The company sells directly to rig",
        "operators and through three regional distributors in Norway, Aberdeen and Houston.",
        "The group employs 214 people across two manufacturing sites.",
    ],
    [
        "2. Financial performance",
        "",
        "Revenue for the financial year ended 31 December 2025 was 48.2 million pounds,",
        "up from 41.7 million pounds in the prior year.",
        "",
        "Gross margin was 34.1 per cent in 2025 compared with 32.8 per cent in 2024.",
        "Adjusted EBITDA for 2025 was 9.6 million pounds.",
        "",
        "The top three customers accounted for 62 per cent of revenue in 2025, and the",
        "single largest customer represented 31 per cent of revenue.",
    ],
    [
        "3. Market and competition",
        "",
        "Principal competitors are Bergstrom Fluid Systems and Caldera Couplings.",
        "Bergstrom competes primarily on price in the North Sea market, while Caldera",
        "holds the stronger position in the Gulf of Mexico.",
        "",
        "4. Key risks",
        "",
        "The directors consider the principal risks to be customer concentration,",
        "exposure to the oil price cycle, and the pending renewal of the Houston site",
        "lease which expires in March 2027.",
    ],
]

DOCX_SECTIONS = [
    ("Management meeting notes", [
        "Notes taken at the management presentation held on 14 January 2026.",
        "Present: the chief executive, the finance director and two advisers from INVRT.",
    ]),
    ("Operations", [
        "Management confirmed that the Aberdeen site runs at approximately 78 per cent of",
        "capacity and that the Houston site runs at approximately 64 per cent of capacity.",
        "A capital expenditure programme of 3.4 million pounds is planned for 2026.",
    ]),
    ("Commercial", [
        "Management stated that the contract with the largest customer runs to December 2027",
        "and contains a two year extension option exercisable by the customer.",
        "No pricing concessions were given in the 2025 renewal round.",
    ]),
]

DOCX_TABLE = [
    ["Site", "Headcount", "Capacity utilisation"],
    ["Aberdeen", "126", "78 per cent"],
    ["Houston", "88", "64 per cent"],
]


def escape(text: str) -> str:
    return text.replace("\\", r"\\").replace("(", r"\(").replace(")", r"\)")


def content_stream(lines) -> bytes:
    parts = ["BT", "/F1 11 Tf", "54 740 Td", "15 TL"]
    for line in lines:
        parts.append(f"({escape(line)}) Tj")
        parts.append("T*")
    parts.append("ET")
    return "\n".join(parts).encode("latin-1", "replace")


def build_pdf(path: str):
    objects = {}
    page_ids = [3 + i * 2 for i in range(len(PAGES))]

    objects[1] = b"<< /Type /Catalog /Pages 2 0 R >>"
    kids = " ".join(f"{pid} 0 R" for pid in page_ids)
    objects[2] = f"<< /Type /Pages /Kids [{kids}] /Count {len(PAGES)} >>".encode()

    font_id = 3 + len(PAGES) * 2
    for index, lines in enumerate(PAGES):
        page_id = page_ids[index]
        content_id = page_id + 1
        objects[page_id] = (
            f"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] "
            f"/Contents {content_id} 0 R /Resources << /Font << /F1 {font_id} 0 R >> >> >>"
        ).encode()
        stream = content_stream(lines)
        objects[content_id] = b"<< /Length " + str(len(stream)).encode() + b" >>\nstream\n" + stream + b"\nendstream"
    objects[font_id] = b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>"

    out = bytearray(b"%PDF-1.4\n")
    offsets = {}
    for obj_id in sorted(objects):
        offsets[obj_id] = len(out)
        out += f"{obj_id} 0 obj\n".encode() + objects[obj_id] + b"\nendobj\n"

    xref_start = len(out)
    count = max(objects) + 1
    out += f"xref\n0 {count}\n".encode()
    out += b"0000000000 65535 f \n"
    for obj_id in range(1, count):
        out += f"{offsets.get(obj_id, 0):010d} 00000 n \n".encode()
    out += f"trailer\n<< /Size {count} /Root 1 0 R >>\nstartxref\n{xref_start}\n%%EOF\n".encode()

    with open(path, "wb") as handle:
        handle.write(bytes(out))


def build_docx(path: str):
    from docx import Document

    document = Document()
    document.add_heading("Acme Industrial Holdings — diligence notes", level=0)
    for heading, paragraphs in DOCX_SECTIONS:
        document.add_heading(heading, level=1)
        for paragraph in paragraphs:
            document.add_paragraph(paragraph)

    document.add_heading("Headcount by site", level=1)
    table = document.add_table(rows=0, cols=3)
    table.style = "Table Grid"
    for row in DOCX_TABLE:
        cells = table.add_row().cells
        for cell, value in zip(cells, row):
            cell.text = value

    document.save(path)


if __name__ == "__main__":
    target = sys.argv[1] if len(sys.argv) > 1 else "fixtures"
    os.makedirs(target, exist_ok=True)
    pdf_path = os.path.join(target, "acme-cim.pdf")
    docx_path = os.path.join(target, "acme-notes.docx")
    build_pdf(pdf_path)
    build_docx(docx_path)
    print(f"wrote {pdf_path}")
    print(f"wrote {docx_path}")
