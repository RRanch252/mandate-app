"""Render an approved briefing memo to Word.

Reads a JSON job on stdin, writes the .docx to job["out_path"], and prints a JSON
result on stdout.

Every claim in the body carries a bracketed marker that resolves in the citation
appendix at the end of the document, so the memo can be checked against its sources
after it leaves the app. The AI-assisted disclosure is written into the document
itself rather than only shown in the UI, because the file is the thing that gets
forwarded to a client.
"""

import json
import sys

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

DISCLOSURE = (
    "AI-assisted, human-reviewed. Every statement below was drafted from documents "
    "uploaded to this mandate and carries a citation to its source. Statements marked "
    "\"Not in corpus\" could not be supported by the uploaded documents and no figure "
    "has been supplied for them. Statements marked \"Unverified\" have not been checked "
    "by a reviewer. This memo must be read against its sources before it is relied on."
)

STATUS_LABEL = {
    "filled": "",
    "edited": "Edited by reviewer",
    "unverified": "Unverified",
    "not_in_corpus": "Not in corpus",
    "empty": "Not answered",
    "queued": "Still running",
}


def add_meta(document, label, value):
    paragraph = document.add_paragraph()
    run = paragraph.add_run(f"{label}: ")
    run.bold = True
    paragraph.add_run(str(value))
    paragraph.paragraph_format.space_after = Pt(2)


def add_disclosure(document, text, color=(0x8A, 0x50, 0x00)):
    paragraph = document.add_paragraph()
    run = paragraph.add_run(text)
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(*color)
    paragraph.paragraph_format.space_after = Pt(10)
    return paragraph


def markers_text(markers):
    return " ".join(f"[{m}]" for m in markers) if markers else ""


def build(job):
    memo = job["memo"]
    document = Document()

    title = document.add_heading(memo.get("title") or "Briefing memo", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT

    add_meta(document, "Mandate", memo.get("mandate_name", ""))
    if memo.get("client_label"):
        add_meta(document, "Client", memo["client_label"])
    add_meta(document, "Prepared by", memo.get("generated_by", ""))
    add_meta(document, "Approved by", memo.get("approved_by") or "Pending approval")
    add_meta(document, "Generated", memo.get("generated_at", ""))
    add_meta(document, "Answer mode", memo.get("model", "extractive"))

    add_disclosure(document, DISCLOSURE)

    sections = memo.get("sections") or []
    if sections:
        document.add_heading("Evidence summary", level=1)
    for section in sections:
        document.add_heading(section.get("heading") or "Topic", level=2)
        rows = section.get("rows") or []
        if not rows:
            document.add_paragraph("No questions answered for this topic.")
            continue
        table = document.add_table(rows=1, cols=3)
        table.style = "Light Grid Accent 1"
        header = table.rows[0].cells
        header[0].text = "Question"
        header[1].text = "Finding"
        header[2].text = "Sources"
        for cell in header:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True
        for row in rows:
            cells = table.add_row().cells
            cells[0].text = row.get("question", "")
            status = row.get("status", "filled")
            body = (row.get("answer") or "").strip()
            label = STATUS_LABEL.get(status, "")
            if status in ("not_in_corpus", "empty", "queued"):
                # Deliberately no number, no paraphrase, no hedge: the memo states
                # that the documents do not answer this.
                cells[1].text = "Not stated in the uploaded documents."
            else:
                cells[1].text = body
            if label and status not in ("not_in_corpus",):
                note = cells[1].add_paragraph()
                run = note.add_run(label)
                run.italic = True
                run.font.size = Pt(8)
            cells[2].text = markers_text(row.get("markers"))

    qa = memo.get("qa") or []
    if qa:
        document.add_heading("Questions asked", level=1)
        for item in qa:
            question = document.add_paragraph()
            run = question.add_run(item.get("question", ""))
            run.bold = True
            answer_text = (item.get("answer") or "").strip()
            if item.get("not_in_corpus"):
                answer_text = "Not stated in the uploaded documents."
            body = document.add_paragraph(answer_text)
            marks = markers_text(item.get("markers"))
            if marks:
                body.add_run(f" {marks}")

    appendix = memo.get("appendix") or []
    document.add_heading("Citation appendix", level=1)
    if not appendix:
        document.add_paragraph("No citations were attached to this memo.")
    for entry in appendix:
        paragraph = document.add_paragraph()
        marker = paragraph.add_run(f"[{entry.get('marker')}] ")
        marker.bold = True
        source = paragraph.add_run(f"{entry.get('document', '')} — {entry.get('locator', '')}")
        source.bold = True
        quote_paragraph = document.add_paragraph()
        quote_run = quote_paragraph.add_run(f"\u201c{(entry.get('quote') or '').strip()}\u201d")
        quote_run.italic = True
        quote_run.font.size = Pt(9)
        quote_paragraph.paragraph_format.left_indent = Pt(18)
        quote_paragraph.paragraph_format.space_after = Pt(8)

    document.save(job["out_path"])


def main():
    try:
        job = json.load(sys.stdin)
        build(job)
        json.dump({"ok": True, "path": job["out_path"]}, sys.stdout)
    except Exception as exc:
        json.dump({"ok": False, "error": str(exc)}, sys.stdout)
    sys.stdout.flush()


if __name__ == "__main__":
    main()
