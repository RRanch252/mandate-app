"""Parse one PDF or DOCX into citable chunks.

Reads a JSON job on stdin, writes a JSON result on stdout. Nothing else may be
written to stdout, because the Node API parses it directly; diagnostics go to stderr.

A chunk is only useful if a human can be taken back to the exact place it came from,
so every chunk carries a locator: a page number for PDFs, a heading plus paragraph
ordinal for DOCX. Chunks that cannot be located are dropped rather than indexed,
because an uncitable chunk would let the answerer make a claim it cannot support.

Job:    {"path": "...", "ext": "pdf"|"docx"}
Result: {"ok": true, "page_count": int|null, "chunks": [...]}
        {"ok": false, "error": "..."}
"""

import json
import re
import sys

# Chunks are sized to hold a full argument (a paragraph or two) while staying small
# enough that a citation points at something a reader can scan in a few seconds.
TARGET_CHARS = 1100
MIN_CHARS = 60

SENTENCE_SPLIT = re.compile(r"(?<=[.!?])\s+(?=[A-Z(\"'\u201c])")


def reflow(text: str) -> str:
    """Rejoin lines that a PDF extractor broke mid-sentence.

    PDF text extraction returns one line per rendered line, so a single sentence
    arrives split across several lines. Left alone this shreds sentences, and a
    citation that quotes half a sentence is worse than no citation.

    A line is treated as a continuation only if the line above it looks wrapped: it
    ran near the full measure, or this line starts lower case. Without that test a
    short heading gets glued onto the paragraph beneath it, which makes the resulting
    sentence longer and pushes it down the ranking against its own body text.
    """
    wrapped_line_min = 45
    out: list[str] = []
    for raw in text.split("\n"):
        line = raw.strip()
        if not line:
            out.append("")
            continue
        previous = out[-1] if out else ""
        continues = bool(previous) and not re.search(r"[.!?:;]$", previous) and (
            len(previous) >= wrapped_line_min or line[:1].islower()
        )
        if continues:
            if previous.endswith("-"):
                out[-1] = previous[:-1] + line  # word hyphenated across a line break
            else:
                out[-1] = f"{previous} {line}"
        else:
            out.append(line)
    return "\n".join(out)


def normalise(text: str) -> str:
    text = text.replace("\u00a0", " ")
    text = re.sub(r"[ \t]+", " ", text)
    text = reflow(text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def split_to_target(text: str):
    """Split text into pieces of roughly TARGET_CHARS, preferring sentence boundaries."""
    text = normalise(text)
    if not text:
        return []
    if len(text) <= TARGET_CHARS:
        return [text]

    pieces = []
    current = ""
    for sentence in SENTENCE_SPLIT.split(text):
        # A single sentence longer than the target is hard-wrapped on whitespace so
        # that pathological documents (tables flattened into one line) still index.
        while len(sentence) > TARGET_CHARS:
            cut = sentence.rfind(" ", 0, TARGET_CHARS)
            if cut <= 0:
                cut = TARGET_CHARS
            if current:
                pieces.append(current.strip())
                current = ""
            pieces.append(sentence[:cut].strip())
            sentence = sentence[cut:].lstrip()
        if len(current) + len(sentence) + 1 > TARGET_CHARS and current:
            pieces.append(current.strip())
            current = sentence
        else:
            current = f"{current} {sentence}".strip() if current else sentence
    if current.strip():
        pieces.append(current.strip())
    return [p for p in pieces if len(p) >= MIN_CHARS or len(pieces) == 1]


def parse_pdf(path: str):
    from pypdf import PdfReader

    reader = PdfReader(path)
    if reader.is_encrypted:
        # An empty-password decrypt covers documents that are protected only against
        # editing, which is common for CIMs. A real password is a hard failure.
        try:
            reader.decrypt("")
        except Exception:
            raise RuntimeError("PDF is password protected")

    chunks = []
    ordinal = 0
    for page_number, page in enumerate(reader.pages, start=1):
        try:
            raw = page.extract_text() or ""
        except Exception as exc:  # a single unreadable page must not lose the document
            print(f"page {page_number}: {exc}", file=sys.stderr)
            continue
        for piece in split_to_target(raw):
            chunks.append(
                {
                    "ordinal": ordinal,
                    "text": piece,
                    "locator_kind": "page",
                    "page": page_number,
                    "section": None,
                    "para_index": None,
                }
            )
            ordinal += 1

    if not chunks:
        raise RuntimeError(
            "no extractable text — the PDF is probably a scan and needs OCR, which this MVP does not do"
        )
    return len(reader.pages), chunks


def parse_docx(path: str):
    from docx import Document

    document = Document(path)
    chunks = []
    ordinal = 0
    section = "Body"
    buffer = []
    buffer_start = 0

    def flush():
        nonlocal buffer, ordinal
        if not buffer:
            return
        joined = "\n".join(buffer)
        for piece in split_to_target(joined):
            chunks.append(
                {
                    "ordinal": ordinal,
                    "text": piece,
                    "locator_kind": "section",
                    "page": None,
                    "section": section,
                    "para_index": buffer_start,
                }
            )
            ordinal += 1
        buffer = []

    for index, paragraph in enumerate(document.paragraphs):
        text = (paragraph.text or "").strip()
        style = (paragraph.style.name or "") if paragraph.style is not None else ""
        if style.startswith("Heading") or style == "Title":
            flush()
            if text:
                section = text
            continue
        if not text:
            continue
        if not buffer:
            buffer_start = index
        buffer.append(text)
        if sum(len(line) for line in buffer) >= TARGET_CHARS:
            flush()
    flush()

    # Word tables hold most of the numbers in a diligence pack, so they are indexed
    # as their own chunks rather than being silently skipped.
    for table_index, table in enumerate(document.tables, start=1):
        rows = []
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            if any(cells):
                rows.append(" | ".join(cells))
        if not rows:
            continue
        for piece in split_to_target("\n".join(rows)):
            chunks.append(
                {
                    "ordinal": ordinal,
                    "text": piece,
                    "locator_kind": "section",
                    "page": None,
                    "section": f"Table {table_index}",
                    "para_index": None,
                }
            )
            ordinal += 1

    if not chunks:
        raise RuntimeError("no extractable text in the document")
    return None, chunks


def main():
    try:
        job = json.load(sys.stdin)
        path = job["path"]
        ext = job["ext"].lower()
        if ext == "pdf":
            page_count, chunks = parse_pdf(path)
        elif ext == "docx":
            page_count, chunks = parse_docx(path)
        else:
            raise RuntimeError(f"unsupported extension: {ext}")
        json.dump({"ok": True, "page_count": page_count, "chunks": chunks}, sys.stdout)
    except Exception as exc:
        json.dump({"ok": False, "error": str(exc)}, sys.stdout)
    sys.stdout.flush()


if __name__ == "__main__":
    main()
